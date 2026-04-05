"""DOCX reader — extracts text with heading-based chunking via python-docx."""

from __future__ import annotations

from pathlib import Path

from uofa_cli.document_reader import DocumentChunk

_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}


def read_docx(path: Path) -> list[DocumentChunk]:
    """Read a DOCX and return chunks split by heading sections."""
    import docx

    doc = docx.Document(path)

    sections: list[tuple[str | None, list[str]]] = []
    current_heading: str | None = None
    current_paragraphs: list[str] = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name in _HEADING_STYLES:
            # Save previous section
            if current_paragraphs:
                sections.append((current_heading, current_paragraphs))
            current_heading = para.text.strip()
            current_paragraphs = []
        else:
            text = para.text.strip()
            if text:
                current_paragraphs.append(text)

    # Save last section
    if current_paragraphs:
        sections.append((current_heading, current_paragraphs))

    # If no sections were created, treat entire doc as one chunk
    if not sections:
        full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        return [DocumentChunk(
            text=full_text or "(empty document)",
            source_file=path.name,
            source_path=str(path),
            format="docx",
        )]

    chunks: list[DocumentChunk] = []
    for heading, paragraphs in sections:
        text = "\n".join(paragraphs)
        if heading:
            text = f"## {heading}\n\n{text}"
        chunks.append(DocumentChunk(
            text=text,
            source_file=path.name,
            source_path=str(path),
            section_heading=heading,
            format="docx",
        ))

    return chunks
